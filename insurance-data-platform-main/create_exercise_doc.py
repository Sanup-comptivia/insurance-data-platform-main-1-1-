"""
Generate the trainee hands-on exercise Word document.
Run: python create_exercise_doc.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ───────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size, color in [("Heading 1", 22, "1B4F72"), ("Heading 2", 16, "2E75B6"), ("Heading 3", 13, "2E75B6")]:
    s = doc.styles[level]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True

# ── Helper functions ─────────────────────────────────────────────────────
def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)
    # Light gray background via shading
    shading = run._element.get_or_add_rPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F2F2F2"
    })
    shading.append(sh)
    return p

def add_info_box(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Insurance Data Platform")
run.font.size = Pt(32)
run.font.color.rgb = RGBColor.from_string("1B4F72")
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hands-On Exercise Workbook")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor.from_string("2E75B6")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Data Engineering Training Program")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comptivia Inc.")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
doc.add_paragraph()

add_table(
    ["Detail", "Value"],
    [
        ["Technology Stack", "PySpark, Delta Lake, Databricks, GitHub Actions"],
        ["Architecture", "Medallion (Bronze \u2192 Silver \u2192 Gold)"],
        ["Repository", "insurance-data-platform"],
        ["Difficulty Levels", "Beginner \u2192 Intermediate \u2192 Advanced"],
        ["Total Exercises", "10 (across 5 modules)"],
    ],
    col_widths=[2.0, 4.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "Introduction & Prerequisites", "3"),
    ("2.", "Module 1 \u2014 Bronze Layer: Add a New Data Source", "4"),
    ("3.", "Module 2 \u2014 Silver Layer: Build a New Transformer", "7"),
    ("4.", "Module 3 \u2014 Gold Layer: Create a New Dimension & Extend a Fact", "10"),
    ("5.", "Module 4 \u2014 Data Quality: Add Custom Validation Rules", "14"),
    ("6.", "Module 5 \u2014 End-to-End: Integrate a New Line of Business", "17"),
    ("7.", "Appendix A \u2014 Evaluation Rubric", "20"),
    ("8.", "Appendix B \u2014 Reference Commands", "21"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(11)
    run2 = p.add_run(f"  ........  {page}")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction & Prerequisites", level=1)

doc.add_paragraph(
    "This workbook contains hands-on exercises designed to give you practical experience "
    "modifying a production-style data engineering platform. Each exercise requires you to "
    "make real code changes in the insurance-data-platform repository."
)

doc.add_heading("Learning Objectives", level=2)
bullets = [
    "Understand the Medallion Architecture (Bronze \u2192 Silver \u2192 Gold) by working within each layer.",
    "Write PySpark transformations that cast types, deduplicate, and validate data.",
    "Design and implement star schema dimensions and fact tables.",
    "Build data quality checks using the DataQualityChecker framework.",
    "Integrate a brand-new line of business end-to-end through the pipeline.",
    "Follow production patterns: idempotent loads, surrogate keys, metadata tracking.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Prerequisites", level=2)
doc.add_paragraph(
    "Before starting these exercises, ensure you have the following:"
)
prereqs = [
    "Python 3.10+ installed with pip.",
    'The repository cloned: git clone <repo-url> && cd insurance-data-platform.',
    "Dependencies installed: pip install -r requirements.txt.",
    "A code editor (VS Code recommended with Python and Databricks extensions).",
    "Basic familiarity with Python, SQL, and PySpark (column operations, DataFrame API).",
    "Read the README.md and browse the src/ directory structure before starting.",
]
for p_text in prereqs:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_heading("How to Use This Workbook", level=2)
doc.add_paragraph(
    "Each module contains 2 exercises that build on each other. Exercises are graded "
    "by difficulty:"
)
add_table(
    ["Level", "Icon", "Description"],
    [
        ["Beginner", "\u2B50", "Follow the pattern of existing code. Mostly copy-and-adapt."],
        ["Intermediate", "\u2B50\u2B50", "Requires understanding the pattern and making design decisions."],
        ["Advanced", "\u2B50\u2B50\u2B50", "Open-ended. Requires architectural thinking and cross-layer integration."],
    ],
    col_widths=[1.5, 0.8, 4.2],
)

doc.add_paragraph()
add_info_box("Tip", "Read the TRAINEE GUIDE and TRAINEE NOTE comments throughout the codebase. They explain the WHY behind every pattern.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 1 — BRONZE LAYER
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Module 1 \u2014 Bronze Layer: Add a New Data Source", level=1)
doc.add_paragraph(
    "The Bronze layer ingests raw data from external sources into Delta tables without "
    "transforming it. In this module you will add a new data source to the ingestion pipeline."
)

# Exercise 1.1
doc.add_heading("Exercise 1.1 \u2014 Add a New CSV Schema for Professional Liability (E&O)", level=2)
add_info_box("Difficulty", "\u2B50 Beginner")
add_info_box("Files to Modify", "src/common/schemas.py")
add_info_box("Time Estimate", "30\u201345 minutes")
add_info_box("Concepts Practiced", "PySpark StructType/StructField, schema-on-read, Bronze conventions")

doc.add_paragraph()
doc.add_heading("Background", level=3)
doc.add_paragraph(
    "Professional Liability insurance (also called Errors & Omissions or E&O) protects "
    "professionals (doctors, lawyers, accountants, IT consultants) against claims of "
    "negligence or inadequate work. Your task is to add a Bronze CSV schema for this "
    "new line of business."
)

doc.add_heading("Requirements", level=3)
doc.add_paragraph("Add a new schema PROFESSIONAL_LIABILITY_CSV to the BronzeSchemas class with these columns:")

add_table(
    ["Column Name", "Type", "Description"],
    [
        ["policy_id", "StringType()", "Unique policy identifier (PK)"],
        ["policy_number", "StringType()", "Human-readable policy number"],
        ["effective_date", "StringType()", "Policy start date (raw string in Bronze)"],
        ["expiration_date", "StringType()", "Policy end date"],
        ["insured_name", "StringType()", "Name of the professional or firm"],
        ["profession_type", "StringType()", "E.g., Lawyer, Doctor, Accountant, IT Consultant"],
        ["license_number", "StringType()", "Professional license number"],
        ["state_code", "StringType()", "State where the professional is licensed"],
        ["retroactive_date", "StringType()", "Earliest date claims can be reported for"],
        ["per_claim_limit", "StringType()", "Max payout per individual claim"],
        ["aggregate_limit", "StringType()", "Max payout per policy period"],
        ["deductible_amount", "StringType()", "Amount insured pays before coverage kicks in"],
        ["premium_amount", "StringType()", "Annual premium"],
        ["claim_id", "StringType()", "Claim identifier (if a claim exists)"],
        ["date_of_occurrence", "StringType()", "When the alleged error/omission occurred"],
        ["claim_description", "StringType()", "Description of the alleged negligence"],
        ["claimant_name", "StringType()", "Name of the person making the claim"],
        ["claimed_amount", "StringType()", "Amount the claimant is seeking"],
        ["paid_amount", "StringType()", "Amount actually paid on the claim"],
        ["claim_status", "StringType()", "Open, Closed, Pending, Denied"],
        ["agent_id", "StringType()", "Agent who sold the policy"],
        ["agent_name", "StringType()", "Agent's full name"],
        ["underwriter", "StringType()", "Underwriter who approved the policy"],
    ],
    col_widths=[1.8, 1.2, 3.5],
)

doc.add_paragraph()
doc.add_heading("Steps", level=3)
steps = [
    "Open src/common/schemas.py and study the existing BronzeSchemas class.",
    "Notice that ALL Bronze CSV columns are StringType() \u2014 this is intentional (schema-on-read).",
    "Add a new class attribute PROFESSIONAL_LIABILITY_CSV following the same pattern as GENERAL_LIABILITY_CSV.",
    "Add a section comment header (# Professional Liability) following the existing convention.",
    'Run the schema tests to verify: pytest tests/test_schemas.py -v.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Validation Checklist", level=3)
checks = [
    "All columns are StringType() (Bronze convention).",
    "All columns have nullable=True (3rd argument).",
    "The schema is a class attribute (not an instance attribute).",
    "Section comment header follows the existing pattern.",
    "Tests pass: pytest tests/test_schemas.py.",
]
for c in checks:
    doc.add_paragraph(c, style="List Bullet")

doc.add_page_break()

# Exercise 1.2
doc.add_heading("Exercise 1.2 \u2014 Ingest Professional Liability CSV into Bronze Delta", level=2)
add_info_box("Difficulty", "\u2B50\u2B50 Intermediate")
add_info_box("Files to Modify", "src/bronze/ingest_structured.py")
add_info_box("Time Estimate", "30\u201345 minutes")
add_info_box("Concepts Practiced", "Bronze ingestion, Delta table writes, metadata columns, idempotent loads")

doc.add_paragraph()
doc.add_heading("Background", level=3)
doc.add_paragraph(
    "Now that the schema exists, you need to wire up the ingestion pipeline to read "
    "Professional Liability CSV files from the raw landing zone and write them into a "
    "Bronze Delta table."
)

doc.add_heading("Requirements", level=3)
reqs = [
    "Add a new entry to the ingestion_map dictionary in the BronzeStructuredIngestion class.",
    'Source path: {raw_base}/professional_liability/synthetic_policies_claims.',
    'Target table: bronze_professional_liability_csv.',
    'Format: csv.',
    "The ingestion must add metadata columns (_ingestion_timestamp, _source_file, _source_format, _batch_id).",
    "The ingestion must use overwrite mode (idempotent).",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("Steps", level=3)
steps = [
    "Open src/bronze/ingest_structured.py and study the existing ingestion_map.",
    "Locate where other LOBs are listed (property, workers_comp, auto, etc.).",
    "Add a new tuple for Professional Liability following the same pattern.",
    "Verify that the _ingest_source() method will handle your new entry correctly \u2014 trace the code path.",
    "No new methods are needed; the existing generic _ingest_source() handles all CSV sources.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Expected Output", level=3)
doc.add_paragraph("After running the ingestion, you should see:")
add_code_block(
    "  Ingesting: bronze_professional_liability_csv\n"
    "    Source: /mnt/insurance-data/raw/professional_liability/synthetic_policies_claims\n"
    "    Rows: 100,000\n"
    "    Table: insurance_catalog.bronze.bronze_professional_liability_csv"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 2 — SILVER LAYER
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Module 2 \u2014 Silver Layer: Build a New Transformer", level=1)
doc.add_paragraph(
    "The Silver layer cleanses, types, deduplicates, and validates data. Each LOB has "
    "its own transformer class. In this module you will build the transformer for "
    "Professional Liability."
)

# Exercise 2.1
doc.add_heading("Exercise 2.1 \u2014 Define Silver Schemas for Professional Liability", level=2)
add_info_box("Difficulty", "\u2B50 Beginner")
add_info_box("Files to Modify", "src/common/schemas.py")
add_info_box("Time Estimate", "20\u201330 minutes")
add_info_box("Concepts Practiced", "Typed schemas, DecimalType for money, DateType, nullable constraints")

doc.add_paragraph()
doc.add_heading("Requirements", level=3)
doc.add_paragraph(
    "Add two new schemas to the SilverSchemas class: "
    "PROFESSIONAL_LIABILITY_POLICIES and PROFESSIONAL_LIABILITY_CLAIMS."
)

doc.add_paragraph("Policies schema \u2014 key type conversions from Bronze:")
add_table(
    ["Column", "Bronze Type", "Silver Type", "Notes"],
    [
        ["policy_id", "String", "StringType(False)", "NOT nullable \u2014 primary key"],
        ["effective_date", "String", "DateType()", "Cast from string"],
        ["expiration_date", "String", "DateType()", "Cast from string"],
        ["per_claim_limit", "String", "DecimalType(18, 2)", "Money \u2014 always Decimal"],
        ["aggregate_limit", "String", "DecimalType(18, 2)", "Money"],
        ["deductible_amount", "String", "DecimalType(18, 2)", "Money"],
        ["premium_amount", "String", "DecimalType(18, 2)", "Money"],
        ["retroactive_date", "String", "DateType()", "Important for claims-made policies"],
    ],
    col_widths=[1.6, 1.0, 1.8, 2.1],
)

doc.add_paragraph()
doc.add_paragraph("Claims schema \u2014 separate table for claim events:")
add_table(
    ["Column", "Silver Type", "Notes"],
    [
        ["claim_id", "StringType(False)", "NOT nullable \u2014 primary key"],
        ["policy_id", "StringType(True)", "FK to policies table"],
        ["date_of_occurrence", "DateType()", "When the alleged error occurred"],
        ["claimed_amount", "DecimalType(18, 2)", "Amount sought by claimant"],
        ["paid_amount", "DecimalType(18, 2)", "Amount actually paid"],
        ["claim_status", "StringType()", "Open, Closed, Pending, Denied"],
    ],
    col_widths=[1.8, 1.8, 2.9],
)

doc.add_paragraph()
add_info_box("Reminder", "Every Silver schema must include the 3 system columns at the end: _is_valid (BooleanType), _dq_issues (StringType), _ingestion_timestamp (TimestampType).")

doc.add_page_break()

# Exercise 2.2
doc.add_heading("Exercise 2.2 \u2014 Build the ProfessionalLiabilityTransformer Class", level=2)
add_info_box("Difficulty", "\u2B50\u2B50 Intermediate")
add_info_box("Files to Create", "src/silver/transform_professional_liability.py")
add_info_box("Time Estimate", "45\u201360 minutes")
add_info_box("Concepts Practiced", "The 7-step Bronze\u2192Silver pipeline, type casting, DQ checks, deduplication")

doc.add_paragraph()
doc.add_heading("Requirements", level=3)
doc.add_paragraph(
    "Create a new transformer following the established 7-step pattern used by all "
    "existing transformers (study transform_general_liability.py as your reference):"
)

steps_7 = [
    ("Step 1: Read from Bronze", "Read bronze_professional_liability_csv using spark.table()."),
    ("Step 2: Standardize names", "Call standardize_column_names(df) to convert camelCase to snake_case."),
    ("Step 3: Cast types", "Convert strings to proper types: F.to_date() for dates, .cast(DecimalType(18,2)) for money, F.trim() for strings."),
    ("Step 4: Select columns", "Select only the canonical columns defined in your Silver schema."),
    ("Step 5: Deduplicate", "Call deduplicate(df, key_columns=['policy_id']) to keep one row per policy."),
    ("Step 6: DQ checks", "Use DataQualityChecker to validate: policy_id not null, insured_name not null, per_claim_limit positive, aggregate_limit positive, premium_amount positive."),
    ("Step 7: Write to Silver", "Call write_delta_table() with mode='overwrite' to write silver_professional_liability_policies."),
]
for step_name, step_desc in steps_7:
    p = doc.add_paragraph()
    run = p.add_run(f"{step_name}: ")
    run.bold = True
    p.add_run(step_desc)

doc.add_paragraph()
doc.add_heading("Class Structure", level=3)
add_code_block(
    'class ProfessionalLiabilityTransformer:\n'
    '    """Transform professional liability data from Bronze to Silver."""\n'
    '\n'
    '    def __init__(self, spark: SparkSession):\n'
    '        self.spark = spark\n'
    '\n'
    '    def transform_all(self):\n'
    '        self.transform_policies()\n'
    '        self.transform_claims()\n'
    '\n'
    '    def transform_policies(self):\n'
    '        # Implement the 7-step pipeline here\n'
    '        ...\n'
    '\n'
    '    def transform_claims(self):\n'
    '        # Implement the 7-step pipeline here\n'
    '        ...'
)

doc.add_heading("Validation Checklist", level=3)
checks = [
    "Class follows the same structure as GeneralLiabilityTransformer.",
    "All monetary columns use DecimalType(18, 2) \u2014 NOT DoubleType.",
    "Dates are cast using F.to_date(), not manual string parsing.",
    "Deduplication uses the correct key column (policy_id for policies, claim_id for claims).",
    "DQ checks include: not_null on PKs, positive on monetary fields.",
    "System columns (_is_valid, _dq_issues) are added by DataQualityChecker.apply().",
    "Write mode is 'overwrite' (idempotent).",
    "Both transform_policies() and transform_claims() are called from transform_all().",
]
for c in checks:
    doc.add_paragraph(c, style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 3 — GOLD LAYER
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Module 3 \u2014 Gold Layer: Create a New Dimension & Extend a Fact", level=1)
doc.add_paragraph(
    "The Gold layer builds the analytics-ready star schema. Dimensions hold "
    "descriptive attributes; facts hold numeric measures. In this module you will "
    "create a new dimension and extend an existing fact table."
)

# Exercise 3.1
doc.add_heading("Exercise 3.1 \u2014 Create dim_profession Dimension Table", level=2)
add_info_box("Difficulty", "\u2B50\u2B50 Intermediate")
add_info_box("Files to Create", "src/gold/dimensions/dim_profession.py")
add_info_box("Files to Modify", "src/common/schemas.py (add GoldSchemas.DIM_PROFESSION)")
add_info_box("Time Estimate", "45\u201360 minutes")
add_info_box("Concepts Practiced", "Dimension table design, surrogate keys, conformed dimensions")

doc.add_paragraph()
doc.add_heading("Background", level=3)
doc.add_paragraph(
    "Professional Liability policies contain profession-specific attributes that don't "
    "fit naturally into existing dimensions. You need a new dim_profession dimension "
    "to capture the WHO (what type of professional) in the star schema."
)

doc.add_heading("Schema Design", level=3)
add_table(
    ["Column", "Type", "Role", "Description"],
    [
        ["profession_key", "StringType(False)", "PK (surrogate)", "MD5 hash of profession_type + license_state"],
        ["profession_id", "StringType(False)", "NK (natural)", "Derived: profession_type + '_' + state_code"],
        ["profession_type", "StringType()", "Attribute", "Lawyer, Doctor, Accountant, IT Consultant, etc."],
        ["profession_category", "StringType()", "Attribute", "Medical, Legal, Financial, Technology"],
        ["license_state", "StringType()", "Attribute", "State where the professional is licensed"],
        ["requires_board_cert", "BooleanType()", "Attribute", "Whether board certification is required"],
        ["avg_claim_severity", "StringType()", "Attribute", "Low, Medium, High, Very High"],
    ],
    col_widths=[1.6, 1.3, 1.0, 2.6],
)

doc.add_paragraph()
doc.add_heading("Steps", level=3)
steps = [
    "Add the DIM_PROFESSION schema to GoldSchemas in schemas.py.",
    "Create dim_profession.py following the pattern of dim_line_of_business.py (it is a reference/static dimension).",
    "Hardcode the reference data: 8\u201310 profession types with their categories and severity ratings.",
    'Generate surrogate keys using: generate_surrogate_key(df, ["profession_type", "license_state"]).',
    'Write to: insurance_catalog.gold.dim_profession.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Reference Data to Hardcode", level=3)
add_table(
    ["profession_type", "profession_category", "requires_board_cert", "avg_claim_severity"],
    [
        ["Doctor", "Medical", "True", "Very High"],
        ["Surgeon", "Medical", "True", "Very High"],
        ["Dentist", "Medical", "True", "High"],
        ["Lawyer", "Legal", "False", "High"],
        ["Accountant (CPA)", "Financial", "True", "Medium"],
        ["Financial Advisor", "Financial", "False", "Medium"],
        ["IT Consultant", "Technology", "False", "Low"],
        ["Architect", "Engineering", "True", "Medium"],
        ["Real Estate Agent", "Financial", "False", "Low"],
    ],
    col_widths=[1.6, 1.4, 1.5, 1.5],
)

doc.add_page_break()

# Exercise 3.2
doc.add_heading("Exercise 3.2 \u2014 Extend fact_claim_transaction for Professional Liability", level=2)
add_info_box("Difficulty", "\u2B50\u2B50\u2B50 Advanced")
add_info_box("Files to Modify", "src/gold/facts/fact_claim_transaction.py")
add_info_box("Time Estimate", "60\u201390 minutes")
add_info_box("Concepts Practiced", "Fact table extension, FK resolution, union-all-LOBs, metric computation")

doc.add_paragraph()
doc.add_heading("Background", level=3)
doc.add_paragraph(
    "The fact_claim_transaction table currently processes claims from 5 LOBs. You need "
    "to extend it to include Professional Liability claims. This requires adding a new "
    "LOB branch to the union-all-LOBs pattern."
)

doc.add_heading("Requirements", level=3)
reqs = [
    "Add a new branch that reads from silver_professional_liability_claims.",
    "Map Professional Liability claim fields to the fact table's standardized columns.",
    "Resolve FKs: claim_key (from dim_claim), policy_key (from dim_policy), lob_key (from dim_line_of_business), date_key (from dim_date).",
    "Compute metrics: paid_amount, reserve_amount (derive as claimed_amount - paid_amount), recovery_amount (default 0), expense_amount (15% of paid_amount for legal fees).",
    "Calculate total_incurred = paid_amount + reserve_amount + expense_amount - recovery_amount.",
    "Union the new LOB DataFrame with the existing 5 LOB DataFrames.",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("Column Mapping", level=3)
add_table(
    ["Fact Column", "Source (Silver PL Claims)", "Transformation"],
    [
        ["claim_key", "claim_id", "Lookup from dim_claim"],
        ["policy_key", "policy_id", "Lookup from dim_policy"],
        ["date_key", "date_of_occurrence", "Convert to YYYYMMDD integer"],
        ["lob_key", "(hardcoded)", 'Lookup "PL" from dim_line_of_business'],
        ["transaction_type", "(derived)", '"Payment" if paid_amount > 0, else "Reserve"'],
        ["paid_amount", "paid_amount", "Direct mapping"],
        ["reserve_amount", "(derived)", "claimed_amount - paid_amount"],
        ["recovery_amount", "(default)", "lit(0)"],
        ["expense_amount", "(derived)", "paid_amount * 0.15"],
        ["total_incurred", "(derived)", "paid + reserve + expense - recovery"],
        ["is_final_payment", "claim_status", 'True if claim_status == "Closed"'],
    ],
    col_widths=[1.5, 1.8, 3.2],
)

doc.add_paragraph()
add_info_box("Hint", "Study how the existing Auto or General Liability branches are structured in fact_claim_transaction.py. Your new branch follows the exact same pattern \u2014 only the source table name and column mappings differ.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 4 — DATA QUALITY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Module 4 \u2014 Data Quality: Add Custom Validation Rules", level=1)
doc.add_paragraph(
    "Data quality is critical in production pipelines. In this module you will add new "
    "DQ rules and implement a custom validation check."
)

# Exercise 4.1
doc.add_heading("Exercise 4.1 \u2014 Add DQ Rules for Professional Liability", level=2)
add_info_box("Difficulty", "\u2B50 Beginner")
add_info_box("Files to Modify", "config/data_quality_rules.yaml, src/silver/transform_professional_liability.py")
add_info_box("Time Estimate", "20\u201330 minutes")
add_info_box("Concepts Practiced", "DQ rule configuration, check_not_null, check_positive, check_in_set, check_date_range")

doc.add_paragraph()
doc.add_heading("Requirements", level=3)
doc.add_paragraph("Add these DQ rules to the Professional Liability transformer:")

add_table(
    ["Rule", "Check Type", "Rationale"],
    [
        ["policy_id not null", "check_not_null", "Primary key \u2014 must always be present"],
        ["insured_name not null", "check_not_null", "Identifies the professional \u2014 required for dim_insured"],
        ["profession_type in valid set", "check_in_set", 'Must be one of: "Lawyer", "Doctor", "Accountant", "IT Consultant", "Architect", "Financial Advisor", "Dentist", "Surgeon", "Real Estate Agent"'],
        ["per_claim_limit positive", "check_positive", "Coverage limit must be a positive dollar amount"],
        ["aggregate_limit positive", "check_positive", "Annual coverage cap must be positive"],
        ["premium_amount positive", "check_positive", "Premium must be a positive dollar amount"],
        ["effective_date in range", "check_date_range", "Must be between 2000-01-01 and 2030-12-31"],
        ["claim_status in valid set", "check_in_set", 'Must be: "Open", "Closed", "Pending", "Denied"'],
    ],
    col_widths=[2.0, 1.3, 3.2],
)

doc.add_paragraph()
doc.add_heading("Steps", level=3)
steps = [
    "Open config/data_quality_rules.yaml and add a professional_liability section.",
    "List each rule with its check type and parameters (follow the existing LOB format).",
    "Open your ProfessionalLiabilityTransformer and wire the rules into the DQ step.",
    "Use the DataQualityChecker fluent API to chain all checks.",
    'Verify that _is_valid = False and _dq_issues contains the failure reason for bad rows.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_page_break()

# Exercise 4.2
doc.add_heading("Exercise 4.2 \u2014 Implement a Custom Cross-Field Validation", level=2)
add_info_box("Difficulty", "\u2B50\u2B50 Intermediate")
add_info_box("Files to Modify", "src/common/data_quality.py")
add_info_box("Time Estimate", "30\u201345 minutes")
add_info_box("Concepts Practiced", "PySpark column expressions, F.when(), custom DQ check, method chaining")

doc.add_paragraph()
doc.add_heading("Background", level=3)
doc.add_paragraph(
    "The current DataQualityChecker supports single-column checks (not_null, positive, "
    "in_set, etc.). But insurance data often requires CROSS-FIELD validation \u2014 rules "
    "that involve relationships BETWEEN two or more columns."
)

doc.add_heading("Requirements", level=3)
doc.add_paragraph("Add a new method check_less_than_or_equal(col_a, col_b, issue_name) to DataQualityChecker that validates col_a <= col_b.")
doc.add_paragraph()
doc.add_paragraph("Use cases in insurance:")
use_cases = [
    "per_claim_limit <= aggregate_limit (a single claim can't exceed the annual cap)",
    "paid_amount <= claimed_amount (can't pay more than what was claimed)",
    "effective_date <= expiration_date (policy can't expire before it starts)",
    "deductible_amount <= per_claim_limit (deductible can't exceed coverage)",
]
for u in use_cases:
    doc.add_paragraph(u, style="List Bullet")

doc.add_heading("Method Signature", level=3)
add_code_block(
    'def check_less_than_or_equal(self, col_a: str, col_b: str,\n'
    '                              issue_name: str = None) -> "DataQualityChecker":\n'
    '    """\n'
    '    Validate that col_a <= col_b for every row.\n'
    '    Rows where col_a > col_b are flagged as invalid.\n'
    '    \n'
    '    Args:\n'
    '        col_a: The column that should be smaller or equal.\n'
    '        col_b: The column that should be larger or equal.\n'
    '        issue_name: Custom name for the DQ issue (default: "{col_a}_exceeds_{col_b}").\n'
    '    \n'
    '    Returns:\n'
    '        self (for method chaining).\n'
    '    """\n'
    '    # YOUR IMPLEMENTATION HERE\n'
    '    # Hint: Append a tuple to self.checks like existing methods do.\n'
    '    # The condition should be: F.col(col_a) <= F.col(col_b)\n'
    '    ...'
)

doc.add_heading("Validation", level=3)
doc.add_paragraph("After implementing, use it in your ProfessionalLiabilityTransformer:")
add_code_block(
    'dq = DataQualityChecker(df)\n'
    'df = (\n'
    '    dq.check_not_null("policy_id")\n'
    '    .check_positive("per_claim_limit")\n'
    '    .check_positive("aggregate_limit")\n'
    '    .check_less_than_or_equal("per_claim_limit", "aggregate_limit")  # NEW!\n'
    '    .check_less_than_or_equal("deductible_amount", "per_claim_limit")  # NEW!\n'
    '    .apply()\n'
    ')'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 5 — END-TO-END INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Module 5 \u2014 End-to-End: Integrate a New Line of Business", level=1)
doc.add_paragraph(
    "This is the capstone module. You will wire everything together and integrate "
    "Professional Liability as a fully operational 6th line of business."
)

# Exercise 5.1
doc.add_heading("Exercise 5.1 \u2014 Register Professional Liability in Configuration & Dimensions", level=2)
add_info_box("Difficulty", "\u2B50\u2B50 Intermediate")
add_info_box("Files to Modify", "config/pipeline_config.yaml, src/gold/dimensions/dim_line_of_business.py")
add_info_box("Time Estimate", "20\u201330 minutes")
add_info_box("Concepts Practiced", "Configuration management, reference dimension updates, NAIC codes")

doc.add_paragraph()
doc.add_heading("Requirements", level=3)

doc.add_paragraph("1. Add PL to pipeline_config.yaml:")
add_code_block(
    'lines_of_business:\n'
    '  # ... existing 5 LOBs ...\n'
    '  - code: "PL"\n'
    '    name: "Professional Liability"\n'
    '    category: "Specialty"\n'
    '    naic_code: "11.0"\n'
    '    regulatory_body: "State DOI"'
)

doc.add_paragraph()
doc.add_paragraph("2. Add a new row to dim_line_of_business.py for the PL LOB:")
add_code_block(
    '# Add to the reference data list:\n'
    '("PL", "Professional Liability", "Specialty", "State DOI", "11.0")'
)

doc.add_paragraph()
doc.add_paragraph("3. Verify that the surrogate key is generated correctly for the new row.")

doc.add_page_break()

# Exercise 5.2
doc.add_heading("Exercise 5.2 \u2014 Wire Professional Liability into dim_policy and the Pipeline Notebooks", level=2)
add_info_box("Difficulty", "\u2B50\u2B50\u2B50 Advanced")
add_info_box("Files to Modify", "src/gold/dimensions/dim_policy.py, src/gold/dimensions/dim_claim.py, notebooks/04_run_silver.py, notebooks/05_build_gold_dimensions.py")
add_info_box("Time Estimate", "60\u201390 minutes")
add_info_box("Concepts Practiced", "Cross-layer integration, union-all-LOBs pattern, FK resolution, pipeline orchestration")

doc.add_paragraph()
doc.add_heading("Requirements", level=3)
doc.add_paragraph("This exercise ties all previous exercises together into a complete integration:")

doc.add_paragraph()
doc.add_paragraph("Part A \u2014 Extend dim_policy (the central dimension):")
parts_a = [
    "Add a new branch to read from silver_professional_liability_policies.",
    "Map PL policy fields to dim_policy\u2019s standardized columns.",
    "Resolve FKs: lob_key (lookup 'PL'), insured_key, agent_key, location_key, date keys.",
    "Union the PL policies with the existing 5 LOBs.",
]
for p_text in parts_a:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph("Part B \u2014 Extend dim_claim:")
parts_b = [
    "Add a new branch to read from silver_professional_liability_claims.",
    "Map PL claim fields to dim_claim\u2019s standardized columns.",
    "Resolve policy_key FK by joining with dim_policy.",
    "Union with existing LOB claims.",
]
for p_text in parts_b:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph("Part C \u2014 Update pipeline notebooks:")
parts_c = [
    "In 04_run_silver.py: Import and call ProfessionalLiabilityTransformer.",
    "In 05_build_gold_dimensions.py: Ensure dim_profession is built before dim_policy.",
    "Verify the pipeline runs end-to-end without errors.",
]
for p_text in parts_c:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_heading("End-to-End Validation", level=3)
doc.add_paragraph("After completing all exercises, verify:")
validations = [
    "Bronze: bronze_professional_liability_csv table exists with raw data.",
    "Silver: silver_professional_liability_policies and silver_professional_liability_claims exist with typed data.",
    "Gold: dim_profession exists with 9+ profession types.",
    "Gold: dim_policy includes PL policies (query: WHERE lob_key = <PL_key>).",
    "Gold: dim_claim includes PL claims.",
    "Gold: fact_claim_transaction includes PL claim transactions.",
    "DQ: _is_valid flags are correctly set on Silver tables.",
    "Pipeline: run_all_pipeline.py completes without errors.",
]
for v in validations:
    doc.add_paragraph(v, style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX A — EVALUATION RUBRIC
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Appendix A \u2014 Evaluation Rubric", level=1)
doc.add_paragraph(
    "Use this rubric to evaluate trainee submissions. Each exercise is scored on "
    "a 4-point scale."
)

add_table(
    ["Criteria", "4 (Excellent)", "3 (Good)", "2 (Needs Work)", "1 (Incomplete)"],
    [
        [
            "Code Correctness",
            "Code runs without errors; all tests pass.",
            "Minor issues (e.g., wrong column name) but logic is sound.",
            "Significant bugs; partial implementation.",
            "Code does not run or is largely missing.",
        ],
        [
            "Pattern Adherence",
            "Follows existing patterns exactly (naming, structure, conventions).",
            "Mostly follows patterns with minor deviations.",
            "Significant deviations from established patterns.",
            "Does not follow existing patterns.",
        ],
        [
            "Data Types",
            "All types are correct (Decimal for money, Date for dates, etc.).",
            "1\u20132 type errors.",
            "Multiple type errors (e.g., DoubleType for money).",
            "Types not addressed.",
        ],
        [
            "Data Quality",
            "All required DQ checks present; custom checks implemented.",
            "Most DQ checks present.",
            "Some DQ checks missing.",
            "No DQ checks implemented.",
        ],
        [
            "Code Style",
            "Clean, well-commented, follows project conventions.",
            "Readable but could use more comments.",
            "Messy or inconsistent style.",
            "Unreadable or no effort at style.",
        ],
    ],
    col_widths=[1.2, 1.5, 1.5, 1.4, 1.2],
)

doc.add_paragraph()
doc.add_heading("Scoring Summary", level=2)
add_table(
    ["Module", "Exercises", "Max Points"],
    [
        ["Module 1 \u2014 Bronze", "1.1 + 1.2", "40"],
        ["Module 2 \u2014 Silver", "2.1 + 2.2", "40"],
        ["Module 3 \u2014 Gold", "3.1 + 3.2", "40"],
        ["Module 4 \u2014 Data Quality", "4.1 + 4.2", "40"],
        ["Module 5 \u2014 End-to-End", "5.1 + 5.2", "40"],
        ["TOTAL", "", "200"],
    ],
    col_widths=[2.5, 1.5, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX B — REFERENCE COMMANDS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Appendix B \u2014 Reference Commands", level=1)

doc.add_heading("Development Environment", level=2)
add_code_block(
    "# Clone and set up\n"
    "git clone <repo-url> && cd insurance-data-platform\n"
    "python -m venv .venv && source .venv/bin/activate\n"
    "pip install -r requirements.txt\n"
    "\n"
    "# Run tests\n"
    "pytest tests/ -v\n"
    "pytest tests/test_schemas.py -v  # schema tests only\n"
    "pytest tests/test_utils.py -v    # utility tests only\n"
    "\n"
    "# Lint and format\n"
    "flake8 src/ --max-line-length=120\n"
    "black src/ --line-length=120 --check\n"
    "black src/ --line-length=120  # auto-format"
)

doc.add_heading("Git Workflow", level=2)
add_code_block(
    "# Create a feature branch for your exercises\n"
    "git checkout -b trainee/<your-name>/professional-liability\n"
    "\n"
    "# Stage, commit, and push\n"
    "git add src/common/schemas.py\n"
    "git commit -m \"Add Professional Liability Bronze schema\"\n"
    "git push -u origin trainee/<your-name>/professional-liability\n"
    "\n"
    "# Create a PR for review\n"
    "gh pr create --title \"Add Professional Liability LOB\" --body \"...\""
)

doc.add_heading("Useful PySpark Commands (for debugging)", level=2)
add_code_block(
    "# Check a table exists\n"
    "spark.table('insurance_catalog.silver.silver_property_policies').printSchema()\n"
    "\n"
    "# Count rows\n"
    "spark.table('insurance_catalog.silver.silver_property_policies').count()\n"
    "\n"
    "# Check DQ flags\n"
    "df.filter(F.col('_is_valid') == False).select('_dq_issues').show(truncate=False)\n"
    "\n"
    "# Verify surrogate keys\n"
    "df.select('policy_key', 'policy_id').show(5, truncate=False)\n"
    "\n"
    "# Check for duplicates\n"
    "df.groupBy('policy_id').count().filter(F.col('count') > 1).show()"
)

doc.add_heading("Key Imports", level=2)
add_code_block(
    "from pyspark.sql import SparkSession\n"
    "from pyspark.sql import functions as F\n"
    "from pyspark.sql.types import DecimalType, DateType, StringType, BooleanType\n"
    "from src.common.utils import (\n"
    "    standardize_column_names,\n"
    "    deduplicate,\n"
    "    write_delta_table,\n"
    "    generate_surrogate_key,\n"
    "    date_to_key,\n"
    ")\n"
    "from src.common.data_quality import DataQualityChecker"
)

# ── Save ─────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Insurance_Data_Platform_Trainee_Exercises.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
